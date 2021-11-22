import sqlite3
import pandas as pd
import numpy as np
import docx

def cagr(start_value, end_value, num_periods=2):
    return (end_value / start_value) ** (1 / (num_periods - 1)) - 1

conn = sqlite3.connect("test.db")
cursor = conn.cursor()

sql = "SELECT * FROM testidprod WHERE partner is NULL AND state is NULL AND bs=0 AND (factor=1 OR factor=2)"

df = pd.read_sql_query(sql, conn)
df.head()

df1=pd.pivot_table(df,columns=["factor","year"],values=["res"],aggfunc=np.sum,fill_value=np.nan).reset_index(drop=True)#.rename(columns={'res':'1'})
world={0:df1.sum(axis=1)}
df1=df1.rename(index=world)
df1

df2=pd.pivot_table(df,columns=["factor","year"],values=["res"],aggfunc=np.sum,fill_value=np.nan).reset_index(drop=True)#.rename(columns={'res':'1'})
df2=df2[1].join(df2[2],rsuffix='_r', lsuffix='_l').join(df2[2]/df2[1],rsuffix='_r', lsuffix='_l').rename(index=world)
df2.columns = pd.MultiIndex.from_tuples(zip([1]*13+[2]*13+[6]*13, pd.unique(df['year']).tolist()*3),names=['factor', 'year'])
df2

df2.to_excel("report.xlsx")

df6=df2[6].iloc[[0]].to_numpy().tolist()
df6 = pd.DataFrame(df6,columns =[i for i in range(13)])
df6 = df6.transpose()
df6[1]=cagr(df6-df6.diff().fillna(df6),df6)
df6.loc[0,1]=df6.loc[12,1]=np.nan
df6[0]=pd.unique(df['year']).tolist()
df6.columns=['Year', 'World Value']
dfp=pd.DataFrame(df6)
df6.index=pd.MultiIndex.from_tuples(zip([6]*13),names=['Factor'])
df6

doc = docx.Document()
doc.add_heading('Calculating CAGR', 0)
t = doc.add_table(df6.shape[0]+2, df6.shape[1]+1)
for j in range(df6.shape[-1]):
    t.cell(0,j+1).text = df6.columns[j]
for i in range(df6.shape[0]):
    t.cell(i+1,0).text = str(df6.index[i][0])
for i in range(df6.shape[0]):
    for j in range(df6.shape[-1]):
        t.cell(i+1,j+1).text = str(df6.values[i,j])
t.cell(0,0).text = "Factor"
cr=round(cagr(df6.values[1,1],
        df6.values[df6.shape[0]-2,1],
        len(pd.unique(df['year']).tolist())-2),2)
if cr>0:
    factor="grew"
else:
    factor="decreased"
doc.add_paragraph(f'Factor 6 {factor} by avg {cr}% every year from {int(df6.values[1,0])} to {int(df6.values[11,0])}')
doc.save('report.docx')

    #   Возможные варианты решения:
    #
    #   Вар2 - выполнить преобразования столбцов через циклы
    #   Вар3 - через запросы SQL, с последующим доведением в python.